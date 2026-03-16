import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

doc = Document(r'c:\Users\teamp\OneDrive\Documents\Single-Tap\NETWORKING_SCREENS_DOCUMENTATION.docx')

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Deep inspect screenshot card (Table 2)
print("=== SCREENSHOT CARD (Table 2) - FULL XML STRUCTURE ===")
table = doc.tables[2]
cell = table.rows[0].cells[0]

# All paragraphs in the cell
for pi, p in enumerate(cell.paragraphs):
    align = p.alignment
    pf = p.paragraph_format
    runs_detail = []
    for ri, r in enumerate(p.runs):
        text_repr = repr(r.text[:60])
        size = r.font.size
        bold = r.bold
        color = r.font.color.rgb if r.font.color and r.font.color.rgb else 'N/A'
        font = r.font.name
        runs_detail.append(f"  run[{ri}]: text={text_repr} size={size} bold={bold} color={color} font={font}")

    print(f"\npara[{pi}]: align={align} space_before={pf.space_before} space_after={pf.space_after} line_spacing={pf.line_spacing}")
    if not runs_detail:
        print("  (no runs)")
    for rd in runs_detail:
        print(rd)

# Row height for screenshot table
row = table.rows[0]
trPr = row._tr.find(f'{{{W}}}trPr')
if trPr is not None:
    trH = trPr.find(f'{{{W}}}trHeight')
    if trH is not None:
        print(f"\nRow height: val={trH.get(f'{{{W}}}val')} hRule={trH.get(f'{{{W}}}hRule')}")
    else:
        print("\nRow height: auto (no trHeight)")
else:
    print("\nRow height: auto (no trPr)")

# Cell border/shading
tcPr = cell._tc.tcPr
if tcPr is not None:
    shd = tcPr.find(f'{{{W}}}shd')
    if shd is not None:
        print(f"Cell shading: fill={shd.get(f'{{{W}}}fill')} val={shd.get(f'{{{W}}}val')}")
    borders = tcPr.find(f'{{{W}}}tcBorders')
    if borders is not None:
        for side in ['top','bottom','left','right']:
            el = borders.find(f'{{{W}}}{side}')
            if el is not None:
                print(f"Cell border {side}: val={el.get(f'{{{W}}}val')} sz={el.get(f'{{{W}}}sz')} color={el.get(f'{{{W}}}color')}")

# Also check Table 10 (another screenshot card)
print("\n\n=== SCREENSHOT CARD (Table 10) ===")
table10 = doc.tables[10]
cell10 = table10.rows[0].cells[0]
for pi, p in enumerate(cell10.paragraphs):
    align = p.alignment
    runs_detail = []
    for ri, r in enumerate(p.runs):
        text_repr = repr(r.text[:60])
        size = r.font.size
        bold = r.bold
        color = r.font.color.rgb if r.font.color and r.font.color.rgb else 'N/A'
        runs_detail.append(f"  run[{ri}]: text={text_repr} size={size} bold={bold} color={color}")
    print(f"para[{pi}]: align={align}")
    for rd in runs_detail:
        print(rd)

# Check the SCREEN label spacing
print("\n\n=== SCREEN LABEL SPACING ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('SCREEN'):
        pf = p.paragraph_format
        print(f"P[{i}] '{p.text}': space_before={pf.space_before} space_after={pf.space_after}")
        # Check next paragraph (heading)
        if i+1 < len(doc.paragraphs):
            pn = doc.paragraphs[i+1]
            pfn = pn.paragraph_format
            text = pn.text[:40].encode('ascii','replace').decode()
            print(f"  next P[{i+1}] '{text}': space_before={pfn.space_before} space_after={pfn.space_after}")
        break

# File/Purpose line spacing
print("\n=== FILE/PURPOSE SPACING ===")
for i, p in enumerate(doc.paragraphs[:25]):
    text = p.text[:30].encode('ascii','replace').decode()
    if text.startswith('File:') or text.startswith('Purpose:'):
        pf = p.paragraph_format
        print(f"P[{i}] '{text}': space_before={pf.space_before} space_after={pf.space_after}")
