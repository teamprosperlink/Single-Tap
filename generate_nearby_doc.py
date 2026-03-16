import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ═══════════════════════════════════════════════════════════
# PAGE SETUP — match networking doc exactly
# ═══════════════════════════════════════════════════════════
for sec in doc.sections:
    sec.top_margin = 720090
    sec.bottom_margin = 720090
    sec.left_margin = 899795
    sec.right_margin = 899795

# ═══════════════════════════════════════════════════════════
# GLOBAL STYLES
# ═══════════════════════════════════════════════════════════
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)  # 127000 EMU

# H1: 14pt bold #003366, space_before=24pt, space_after=0
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Emu(177800)  # 14pt
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
h1.paragraph_format.space_before = Emu(304800)  # 24pt
h1.paragraph_format.space_after = Emu(0)
h1.paragraph_format.keep_with_next = True
h1.paragraph_format.keep_together = True

# H2: 13pt bold #00508C, space_before=10pt, space_after=0
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Emu(165100)  # 13pt
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x00, 0x50, 0x8C)
h2.paragraph_format.space_before = Emu(127000)  # 10pt
h2.paragraph_format.space_after = Emu(0)
h2.paragraph_format.keep_with_next = True
h2.paragraph_format.keep_together = True

# H3: bold #0064A0, space_before=10pt, space_after=0
h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x00, 0x64, 0xA0)
h3.paragraph_format.space_before = Emu(127000)
h3.paragraph_format.space_after = Emu(0)
h3.paragraph_format.keep_with_next = True
h3.paragraph_format.keep_together = True


# ── Helper functions ──

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)

def screen_label(number):
    """SCREEN X — 10pt, bold, #0066B2, space_after=2pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(25400)  # 2pt
    run = p.add_run(f'SCREEN {number}')
    run.bold = True
    run.font.size = Emu(127000)  # 10pt
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xB2)
    run.font.name = 'Calibri'

def file_line(path):
    """File: path — space_after=4pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(50800)  # 4pt
    r1 = p.add_run('File: ')
    r1.bold = True
    r1.font.size = Emu(127000)
    r1.font.name = 'Calibri'
    r2 = p.add_run(path)
    r2.font.size = Emu(127000)
    r2.font.color.rgb = RGBColor(0x00, 0x66, 0xB2)
    r2.font.name = 'Calibri'

def purpose_line(text):
    """Purpose: text — space_after=8pt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(101600)  # 8pt
    r1 = p.add_run('Purpose: ')
    r1.bold = True
    r1.font.size = Emu(127000)
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Emu(127000)
    r2.font.name = 'Calibri'

def para(text, bold=False, size_emu=127000, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Emu(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Emu(size_emu)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Emu(127000)
    run.font.name = 'Calibri'

def bullet_bold(prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.size = Emu(127000)
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Emu(127000)
    r2.font.name = 'Calibri'

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, color=None, size_emu=127000):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Emu(size_emu)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color

def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    # Remove existing tcW if any
    existing = tcPr.find(f'{{{W_NS}}}tcW')
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

def keep_table_on_one_page(table):
    """Prevent table from splitting across pages."""
    for row in table.rows:
        # cantSplit: prevent a single row from breaking across pages
        trPr = row._tr.get_or_add_trPr()
        cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
        trPr.append(cant_split)
        # keepNext on every paragraph in the row to glue rows together
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._element.get_or_add_pPr()
                keep_next = parse_xml(f'<w:keepNext {nsdecls("w")}/>')
                keep_lines = parse_xml(f'<w:keepLines {nsdecls("w")}/>')
                pPr.append(keep_next)
                pPr.append(keep_lines)

def add_table(headers, rows, col_widths=None):
    """Table Grid with #003366 header, white text."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '003366')
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        if col_widths:
            set_cell_width(cell, col_widths[i])
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_text(cell, val)
            if col_widths:
                set_cell_width(cell, col_widths[ci])
    # Keep table together on one page
    keep_table_on_one_page(table)
    doc.add_paragraph()
    return table

def add_screenshot_placeholder(screen_name):
    """Exact replica of networking doc screenshot card."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]

    # Cell background: #F8F9FA
    set_cell_shading(cell, 'F8F9FA')

    # Cell width: 7937 dxa (full width)
    set_cell_width(cell, 7937)

    # Dashed border: sz=8, color=#AAAAAA
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="dashed" w:sz="8" w:color="AAAAAA" w:space="0"/>'
        f'  <w:bottom w:val="dashed" w:sz="8" w:color="AAAAAA" w:space="0"/>'
        f'  <w:left w:val="dashed" w:sz="8" w:color="AAAAAA" w:space="0"/>'
        f'  <w:right w:val="dashed" w:sz="8" w:color="AAAAAA" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(parse_xml(borders_xml))

    # Para 0: empty + emoji (space_before=60pt=762000, centered, 36pt emoji)
    cell.text = ''
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Emu(762000)  # 60pt top padding
    r0 = p0.add_run('')
    r1 = p0.add_run('\n\n\U0001F4F1')
    r1.font.size = Emu(457200)  # 36pt

    # Para 1: screen name in brackets (14pt, bold, #969696, centered)
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p1.add_run(f'[ {screen_name} ]')
    r2.font.size = Emu(177800)  # 14pt
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x96, 0x96, 0x96)
    r2.font.name = 'Calibri'

    # Para 2: paste instruction (10pt, #B4B4B4, centered)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p2.add_run('Paste screenshot here')
    r3.font.size = Emu(127000)  # 10pt
    r3.font.color.rgb = RGBColor(0xB4, 0xB4, 0xB4)
    r3.font.name = 'Calibri'

    # Keep screenshot card together on one page
    keep_table_on_one_page(table)

    doc.add_paragraph()

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(0.5)


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

# SUPPER APP — 36pt, bold, #003366
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUPPER APP')
run.bold = True
run.font.size = Emu(457200)  # 36pt
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
run.font.name = 'Calibri'

# Nearby Module — 28pt, #0066B2
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nearby Module')
run.font.size = Emu(355600)  # 28pt
run.font.color.rgb = RGBColor(0x00, 0x66, 0xB2)
run.font.name = 'Calibri'

doc.add_paragraph()

# Complete Screen Documentation — 18pt, #646464
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complete Screen Documentation')
run.font.size = Emu(228600)  # 18pt
run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

# Meta table — col widths: 2835, 5669 dxa
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Module', 'Nearby / Marketplace'),
    ('Total Screens', '5 (1 feed + 1 my-posts + 1 create + 1 edit + 1 detail)'),
    ('Stack', 'Flutter, Firebase, Geolocator, Speech-to-Text'),
    ('Last Updated', '03 Mar 2026'),
]
for i, (key, val) in enumerate(meta_data):
    cell_k = meta_table.rows[i].cells[0]
    cell_v = meta_table.rows[i].cells[1]
    set_cell_shading(cell_k, '003366')
    set_cell_text(cell_k, key, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_width(cell_k, 2835)
    set_cell_shading(cell_v, 'F0F5FF')
    set_cell_text(cell_v, val)
    set_cell_width(cell_v, 5669)
keep_table_on_one_page(meta_table)

add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)

toc_data = [
    ('1', 'Nearby Screen (Main Feed)', 'near_by_screen.dart'),
    ('2', 'My Posts Screen', 'near_by_posts _screen.dart'),
    ('3', 'Create Post Screen', 'create_post_screen.dart'),
    ('4', 'Edit Post Screen', 'edit_post_screen.dart'),
    ('5', 'Post Detail Screen', 'near_by_post_detail_screen.dart'),
]
toc_table = doc.add_table(rows=6, cols=3)
toc_table.style = 'Table Grid'
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['#', 'Screen Name', 'File']):
    cell = toc_table.rows[0].cells[i]
    set_cell_shading(cell, '003366')
    set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_width(cell, [850, 4535, 3685][i])
for ri, (num, name, file) in enumerate(toc_data):
    set_cell_text(toc_table.rows[ri+1].cells[0], num)
    set_cell_width(toc_table.rows[ri+1].cells[0], 850)
    set_cell_text(toc_table.rows[ri+1].cells[1], name)
    set_cell_width(toc_table.rows[ri+1].cells[1], 4535)
    set_cell_text(toc_table.rows[ri+1].cells[2], file)
    set_cell_width(toc_table.rows[ri+1].cells[2], 3685)
keep_table_on_one_page(toc_table)

doc.add_paragraph()
para('Additional Sections: Screen Navigation Flow, Firestore Collections, Shared UI Patterns, Key Business Rules',
     size_emu=127000, color=RGBColor(0x64, 0x64, 0x64), align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()


# ═══════════════════════════════════════════════════════════
# SCREEN 1 — NEARBY SCREEN
# ═══════════════════════════════════════════════════════════
screen_label(1)
doc.add_heading('Nearby Screen (Main Feed)', level=1)
file_line('lib/screens/near by/near_by_screen.dart (Bottom Tab Index: 3)')
purpose_line('The primary marketplace feed. Displays posts from other users within a 5 km radius, organized by category tabs. Provides search (text + voice), save/bookmark, and navigation to post details, My Posts, and Create Post.')

add_screenshot_placeholder('Nearby Screen - Main Feed')

doc.add_heading('Constructor Parameters', level=2)
add_table(
    ['Parameter', 'Type', 'Required', 'Description'],
    [['onBack', 'VoidCallback?', 'No', 'Callback when user taps back (used by parent)']],
    col_widths=[1700, 1700, 1000, 4671]
)

doc.add_heading('AppBar Layout', level=2)
add_table(
    ['Element', 'Position', 'Behavior'],
    [
        ['Title "Nearby"', 'Center', 'Poppins bold 20px, white'],
        ['More button (3-dot)', 'Right', 'Navigates to MyPostsScreen'],
        ['TabBar (4 tabs)', 'Bottom', 'Services | Jobs | Products | Donation'],
        ['Background', 'Full width', 'Transparent gradient + white bottom border 0.5px'],
    ],
    col_widths=[2268, 1984, 4819]
)

doc.add_heading('AppBar Styling', level=2)
bullet('Background: Gradient black 40% to black 20% to transparent')
bullet('Bottom border: White 0.5px')
bullet('Transparent elevation, no shadow')

doc.add_heading('Category Tabs (TabController)', level=2)
add_table(
    ['Tab', 'Icon', 'Matching Keywords'],
    [
        ['Services', 'computer_rounded', 'service, repair, cleaning, plumber, electrician, tutor, salon, driver, freelance, help'],
        ['Jobs', 'work_rounded', 'job, hiring, work, vacancy, career, recruit, internship, full time, part time'],
        ['Products', 'shopping_bag_rounded', 'product, sell, buy, sale, price, shop, discount, offer, new, used, or price != null'],
        ['Donation', 'volunteer_activism_rounded', 'isDonation == true or category == "donation"'],
    ],
    col_widths=[1417, 2268, 5386]
)

doc.add_heading('Tab Bar Styling', level=2)
bullet('Indicator: White, full-tab width, 1px weight')
bullet('Selected label: White, Poppins 13px, weight 600')
bullet('Unselected label: White 60% opacity, Poppins 13px, normal weight')
bullet('isScrollable: false (tabs fill width equally)')
bullet('Divider: Transparent')

doc.add_heading('Search Bar', level=2)
add_table(
    ['Feature', 'Details'],
    [
        ['Widget', 'GlassSearchField (glassmorphism style)'],
        ['Hint text', '"Search posts..."'],
        ['Text search', 'Filters by title, description, originalPrompt, userName, hashtags, keywords'],
        ['Voice search', 'speech_to_text package, microphone permission required'],
        ['Locale', 'en_IN (Indian English)'],
        ['Listen duration', 'Up to 30 seconds'],
        ['Pause detection', '3 seconds'],
        ['Silence timer', '5 seconds auto-stop'],
    ],
    col_widths=[2268, 6803]
)

doc.add_heading('Data Loading', level=2)
para('Real-time stream: Subscribes to posts collection ordered by createdAt desc, limit 200.', bold=True, size_emu=127000)

doc.add_heading('Client-side Filtering Pipeline', level=3)
add_table(
    ['Step', 'Filter', 'Description'],
    [
        ['1', 'Dedup by ID', 'Remove duplicate document IDs'],
        ['2', 'Skip dummy', 'isDummyPost == true or userId starts with "dummy_"'],
        ['3', 'One per user', 'Latest post only per userId'],
        ['4', 'Title dedup', 'Same userId + same title = duplicate'],
        ['5', 'Exclude self', 'Hide current user\'s own posts'],
        ['6', 'Active only', 'isActive must be true'],
        ['7', 'Image required', 'Must have at least one image'],
        ['8', 'Donation isolation', 'Donation posts only in Donation tab'],
        ['9', 'Category match', 'Keyword-based category filtering'],
        ['10', 'Search match', 'Text search filter'],
        ['11', '5 km radius', 'Exclude posts beyond 5 km or without GPS'],
    ],
    col_widths=[850, 1984, 6237]
)
bullet('Filtered results cached, invalidated on posts/search/category/location change')
bullet('Pagination: 200 posts per page, triggers load-more 500px before scroll end')

doc.add_heading('Location', level=2)
bullet('Gets GPS on init via Geolocator')
bullet('Falls back to Firestore user profile (users/{uid} latitude/longitude)')
bullet('Real-time location stream updates every 100m of movement')
bullet('Distance calculated with Haversine formula')
bullet('Distance displayed on each card (e.g., "1.2 km" or "350 m")')

doc.add_heading('Post Card Information', level=2)
add_table(
    ['Element', 'Position', 'Details'],
    [
        ['Distance badge', 'Top-left', 'Glassmorphism pill, e.g. "1.2 km"'],
        ['Donation badge', 'Top-left (below distance)', 'Orange pill, if isDonation == true'],
        ['Save button', 'Top-right', 'Bookmark toggle, glassmorphism circle'],
        ['Cover image', 'Full card', 'CachedNetworkImage, BoxFit.cover'],
        ['Username', 'Bottom bar', 'White, Poppins 13px bold'],
        ['Call icon', 'Next to username', 'Green circle, if allowCalls != false'],
        ['Post title', 'Bottom bar', 'White 85% opacity, 12px, max 2 lines'],
        ['Price', 'Bottom bar', 'Green #00D67D, Poppins 13px bold'],
    ],
    col_widths=[1984, 2268, 4819]
)

doc.add_heading('Grid Layout', level=3)
bullet('Card height: 200px')
bullet('2-column masonry grid (SliverMasonryGrid.count)')
bullet('Main-axis spacing: 10, Cross-axis spacing: 10')
bullet('Padding: 12px horizontal, 0 top, 80px bottom')
bullet('Each card wrapped in FloatingCard animation')

doc.add_heading('Floating Action Button (FAB)', level=2)
bullet('Icon: "+" (add), white, 28px')
bullet('Color: AppColors.iosBlue (#007AFF)')
bullet('Shape: Circle')
bullet('Position: Bottom-right, 75px from bottom, 20px from right')
bullet('Action: Opens CreatePostScreen')

doc.add_heading('Access Points', level=2)
add_table(
    ['Entry Point', 'Location', 'Navigation Target'],
    [
        ['More button (3-dot)', 'AppBar right', 'MyPostsScreen'],
        ['+ FAB', 'Bottom-right', 'CreatePostScreen'],
        ['Post card tap', 'Feed grid', 'NearByPostDetailScreen (FlipPageRoute)'],
    ],
    col_widths=[2268, 2835, 3969]
)

doc.add_heading('Body Background', level=2)
para('Gradient: #404040 (top) to #000000 (bottom)', size_emu=127000)

add_page_break()


# ═══════════════════════════════════════════════════════════
# SCREEN 2 — MY POSTS SCREEN
# ═══════════════════════════════════════════════════════════
screen_label(2)
doc.add_heading('My Posts Screen', level=1)
file_line('lib/screens/near by/near_by_posts _screen.dart')
purpose_line('Manages the current user\'s own posts. Three tabs: My Post (active), Saved (bookmarked), Delete (soft-deleted with 30-day auto-delete).')

add_screenshot_placeholder('My Posts Screen - 3 Tabs')

doc.add_heading('AppBar Layout', level=2)
add_table(
    ['Element', 'Position', 'Behavior'],
    [
        ['Back arrow', 'Left', 'Navigator.pop()'],
        ['Title "My Posts"', 'Center', 'Poppins bold 20px, white'],
        ['TabBar (3 tabs)', 'Bottom', 'My Post | Saved | Delete'],
    ],
    col_widths=[2268, 1984, 4819]
)

doc.add_heading('Tab Bar Styling', level=2)
bullet('Indicator: White, full-tab width, 1px weight')
bullet('Selected label: White, Poppins 13px, weight 600')
bullet('Unselected label: White 60% opacity')
bullet('isScrollable: true, tabAlignment: center, label padding: 48px horizontal')

doc.add_heading('Tab 1: My Post', level=2)
para('Firestore Query:', bold=True, size_emu=127000)
add_code_block('posts\n  .where("userId", == currentUserId)\n  .where("isActive", == true)\n  .orderBy("createdAt", descending: true)\n  .limit(50)')

doc.add_heading('Card Actions', level=3)
add_table(
    ['Action', 'Button', 'Behavior'],
    [
        ['Edit', 'edit_outlined icon', 'Opens EditPostScreen(postId, postData)'],
        ['Delete', 'delete_outline icon', 'Soft-delete confirmation dialog'],
    ],
    col_widths=[1417, 2268, 5386]
)

doc.add_heading('Soft Delete Behavior', level=3)
bullet('Sets isActive: false and deletedAt: serverTimestamp()')
bullet('Post moves to Delete tab')
bullet('Snackbar: "Post moved to Delete tab"')

doc.add_heading('Tab 2: Saved', level=2)
para('Firestore Query:', bold=True, size_emu=127000)
add_code_block('users/{uid}/saved_posts\n  .orderBy("savedAt", descending: true)\n  .limit(50)')

doc.add_heading('Card Actions', level=3)
add_table(
    ['Action', 'Behavior'],
    [['Unsave', 'Deletes from saved_posts subcollection, shows snackbar']],
    col_widths=[2268, 6803]
)

doc.add_heading('Tab 3: Delete', level=2)
para('Firestore Query:', bold=True, size_emu=127000)
add_code_block('posts\n  .where("userId", == currentUserId)\n  .where("isActive", == false)\n  .orderBy("createdAt", descending: true)\n  .limit(100)')
bullet('Auto-delete: Posts with deletedAt older than 30 days are permanently deleted')
bullet('Top-left badge: "X days left" — turns red when <= 7 days')

doc.add_heading('Card Actions', level=3)
add_table(
    ['Action', 'Button', 'Behavior'],
    [
        ['Restore', 'restore_rounded', 'Confirmation dialog -> isActive=true, remove deletedAt'],
        ['Delete Forever', 'delete_forever_rounded', 'Confirmation dialog -> permanent Firestore delete'],
    ],
    col_widths=[1700, 2268, 5103]
)

doc.add_heading('Post Card Layout', level=2)
add_table(
    ['Element', 'Details'],
    [
        ['Card height', '240px (taller than feed cards)'],
        ['Layout', '2-column masonry grid, spacing: 10'],
        ['Animation', 'FloatingCard floating animation'],
        ['Bottom bar', 'Glassmorphism — username, title, price'],
        ['Tap action', 'Opens NearByPostDetailScreen (isDeleted flag for deleted)'],
    ],
    col_widths=[2268, 6803]
)

doc.add_heading('Confirmation Dialogs', level=2)
add_table(
    ['Dialog', 'Icon', 'Action Color', 'Action Text'],
    [
        ['Soft Delete', 'delete_outline_rounded', 'AppColors.error (Red)', '"Delete"'],
        ['Restore', 'restore_rounded', 'AppColors.success (Green)', '"Restore"'],
        ['Permanent Delete', 'delete_forever_rounded', 'AppColors.error (Red)', '"Delete Forever"'],
    ],
    col_widths=[1984, 2268, 2268, 2551]
)
bullet('Background: Gradient #404040 to #0F0F0F')
bullet('Border: White 15% opacity, 1px')
bullet('Buttons: Cancel (glass) + Action (colored)')
bullet('Close button (x): Top-right circle')

doc.add_heading('Category Detection Helper', level=2)
add_table(
    ['Condition', 'Returns'],
    [
        ['isDonation == true', '"Donation"'],
        ['category/text contains job, hiring', '"Jobs"'],
        ['category/text contains service, repair, cleaning', '"Services"'],
        ['Default', '"Products"'],
    ],
    col_widths=[5669, 3402]
)

add_page_break()


# ═══════════════════════════════════════════════════════════
# SCREEN 3 — CREATE POST SCREEN
# ═══════════════════════════════════════════════════════════
screen_label(3)
doc.add_heading('Create Post Screen', level=1)
file_line('lib/screens/near by/create_post_screen.dart')
purpose_line('Full-featured form to create a new marketplace post with images, location, price, categories, keywords, speech-to-text, and donation toggle.')

add_screenshot_placeholder('Create Post Screen')

doc.add_heading('Form Fields', level=2)
add_table(
    ['Field', 'Type', 'Required', 'Speech', 'Description'],
    [
        ['Title', 'TextEditingController', 'Yes', 'Yes', 'Post title'],
        ['Description', 'TextEditingController', 'Yes', 'Yes', 'Detailed description'],
        ['Price', 'TextEditingController', 'No', 'No', 'Numeric (hidden when donation)'],
        ['Currency', 'Dropdown', 'No', 'No', 'INR, USD, EUR, GBP, AED, SAR'],
        ['Location', 'TextEditingController', 'No', 'No', 'Auto-detect + search autocomplete'],
        ['Keywords', 'TextEditingController', 'No', 'No', 'Chip-style tags'],
        ['Categories', 'TextEditingController', 'Yes (>=1)', 'No', 'Chip-style category tags'],
        ['Offer', 'TextEditingController', 'No', 'Yes', 'Special offer text'],
        ['Images', 'List<File>', 'Yes (>=1)', '—', 'Up to 3 from camera/gallery'],
    ],
    col_widths=[1417, 1984, 1134, 850, 3685]
)

doc.add_heading('Save Validation (_isFormValid)', level=2)
add_code_block('title.isNotEmpty && description.isNotEmpty\n&& categories.isNotEmpty && images.isNotEmpty')

doc.add_heading('Image Picker Options', level=2)
add_table(
    ['Feature', 'Details'],
    [
        ['Maximum images', '3'],
        ['Sources', 'Camera or Gallery (via image_picker package)'],
        ['Max resolution', '1080 x 1080 pixels'],
        ['JPEG quality', '85%'],
        ['Upload path', 'Firebase Storage: posts/{userId}/post_{timestamp}_{index}.jpg'],
        ['Display', 'Horizontal scrollable thumbnails with remove (x) button'],
    ],
    col_widths=[2268, 6803]
)

doc.add_heading('Location Features', level=2)

doc.add_heading('Auto-detection (on init)', level=3)
bullet('Checks GPS service enabled + permission')
bullet('Gets last known position (instant, from cache)')
bullet('Fallback: getCurrentPosition with 5-second timeout')
bullet('Reverse geocodes to city name via GeocodingService')
bullet('Sets _detectedLat / _detectedLng coordinates')

doc.add_heading('Manual Search with Autocomplete', level=3)
bullet('Debounced (400ms) search as user types (minimum 2 chars)')
bullet('Uses GeocodingService.searchLocation() with user GPS as bias')
bullet('Dropdown suggestions show: area, city, state')
bullet('Selecting suggestion updates coordinates and text field')

doc.add_heading('Speech-to-Text', level=2)
add_table(
    ['Feature', 'Details'],
    [
        ['Per-field mic toggle', 'Each text field has its own mic button'],
        ['Active indicator', '_activeController tracks which field is listening'],
        ['Locale', 'en_IN (Indian English)'],
        ['Listen duration', 'Up to 30 seconds'],
        ['Pause detection', '3 seconds'],
        ['Partial results', 'Shown in real-time as user speaks'],
    ],
    col_widths=[2835, 6236]
)

doc.add_heading('Donation Toggle', level=2)
bullet('When enabled: Hides price field, sets isDonation: true')
bullet('When disabled: Shows price field with currency selector')

doc.add_heading('Currency Options', level=2)
add_table(
    ['Code', 'Symbol', 'Name'],
    [
        ['INR', '\u20b9', 'Indian Rupee (default)'],
        ['USD', '$', 'US Dollar'],
        ['EUR', '\u20ac', 'Euro'],
        ['GBP', '\u00a3', 'British Pound'],
        ['AED', 'AED', 'UAE Dirham'],
        ['SAR', 'SAR', 'Saudi Riyal'],
    ],
    col_widths=[1417, 1417, 6237]
)

doc.add_heading('Allow Calls Toggle', level=2)
bullet('Default: true')
bullet('When enabled: Other users see call icon and can voice-call post creator')

doc.add_heading('Data Flow (on Save)', level=2)
bullet('1. Validate all required fields')
bullet('2. Upload images to Firebase Storage -> get download URLs')
bullet('3. Fetch current user profile from Firestore')
bullet('4. Get GPS coordinates (cached or fresh)')
bullet('5. Build postData map')
bullet('6. Write to Firestore: posts.add(postData)')
bullet('7. Pop screen with result=true')

doc.add_heading('Firestore Document Structure (saved)', level=2)
add_table(
    ['Field', 'Type', 'Example Value'],
    [
        ['title', 'string', '"iPhone 13 Pro Max"'],
        ['description', 'string', '"Brand new, sealed pack..."'],
        ['originalPrompt', 'string', '"iPhone 13 Pro Max"'],
        ['userId', 'string', '"abc123"'],
        ['userName', 'string', '"John Doe"'],
        ['userPhoto', 'string', '"https://..."'],
        ['isActive', 'boolean', 'true'],
        ['createdAt', 'Timestamp', 'serverTimestamp()'],
        ['updatedAt', 'Timestamp', 'serverTimestamp()'],
        ['expiresAt', 'Timestamp', 'now + 30 days'],
        ['allowCalls', 'boolean', 'true'],
        ['isDonation', 'boolean', 'false'],
        ['currency', 'string', '"INR"'],
        ['category', 'string', '"Electronics" (first)'],
        ['categories', 'string[]', '["Electronics", "Gadgets"]'],
        ['keywords', 'string[]', '["iphone", "apple"]'],
        ['location', 'string', '"Mumbai"'],
        ['offer', 'string', '"10% off"'],
        ['hashtags', 'string[]', '[]'],
        ['latitude', 'number', '19.0760'],
        ['longitude', 'number', '72.8777'],
        ['imageUrl', 'string', '"https://..." (first image)'],
        ['images', 'string[]', '["https://...", ...]'],
        ['price', 'number (opt)', '85000'],
    ],
    col_widths=[1984, 1984, 5103]
)

add_page_break()


# ═══════════════════════════════════════════════════════════
# SCREEN 4 — EDIT POST SCREEN
# ═══════════════════════════════════════════════════════════
screen_label(4)
doc.add_heading('Edit Post Screen', level=1)
file_line('lib/screens/near by/edit_post_screen.dart')
purpose_line('Pre-populated form to edit an existing post. Mirrors Create Post Screen but loads existing data and handles image replacement.')

add_screenshot_placeholder('Edit Post Screen')

doc.add_heading('Constructor Parameters', level=2)
add_table(
    ['Parameter', 'Type', 'Required', 'Description'],
    [
        ['postId', 'String', 'Yes', 'Firestore document ID of the post to edit'],
        ['postData', 'Map<String, dynamic>', 'Yes', 'Current post data (pre-populates all form fields)'],
    ],
    col_widths=[1700, 2551, 1000, 3820]
)

doc.add_heading('Key Differences from Create Screen', level=2)
add_table(
    ['Feature', 'Create Screen', 'Edit Screen'],
    [
        ['Form fields', 'Empty on load', 'Pre-populated from postData'],
        ['Images', 'Only new local images', 'Existing URLs + new local images'],
        ['Image limit', '3 new images', '3 total (existing + new)'],
        ['Submit button', '"Create Post"', '"Update Post"'],
        ['Firestore op', '.add(postData)', '.doc(postId).update(postData)'],
        ['Expiry reset', 'Set to now + 30 days', 'Reset to now + 30 days'],
    ],
    col_widths=[1984, 3402, 3685]
)

doc.add_heading('Pre-populated Fields (in initState)', level=2)
add_table(
    ['Field', 'Source Key', 'Default'],
    [
        ['Title', 'postData["title"]', '""'],
        ['Description', 'postData["description"]', '""'],
        ['Price', 'postData["price"]', '(empty)'],
        ['Location', 'postData["location"]', '""'],
        ['Offer', 'postData["offer"]', '""'],
        ['Allow Calls', 'postData["allowCalls"]', 'true'],
        ['Is Donation', 'postData["isDonation"]', 'false'],
        ['Currency', 'postData["currency"]', '"INR"'],
        ['Categories', 'postData["categories"] or ["category"]', '[]'],
        ['Keywords', 'postData["keywords"]', '[]'],
        ['Coordinates', 'postData["latitude"]/["longitude"]', 'null'],
        ['Existing Images', 'postData["imageUrl"] + ["images"]', '[]'],
    ],
    col_widths=[1984, 3969, 3118]
)

doc.add_heading('Image Handling', level=2)
bullet('Displays existing images (URLs) alongside new local images')
bullet('Existing images can be removed individually (x button)')
bullet('New images uploaded to Firebase Storage, combined with remaining existing URLs')
bullet('Total limit: 3 images (existing + new combined)')

doc.add_heading('All Other Features', level=2)
para('Identical to Create Post Screen: Speech-to-text per field, location auto-detect + search autocomplete, donation toggle, currency selector, allow calls toggle, same glassmorphism UI theme.', size_emu=127000)

doc.add_heading('Data Flow (on Save)', level=2)
bullet('1. Validate all required fields')
bullet('2. Upload NEW images to Firebase Storage')
bullet('3. Merge existing image URLs + new upload URLs')
bullet('4. Get GPS coordinates (cached from postData or fresh)')
bullet('5. Build update map')
bullet('6. Write to Firestore: posts.doc(postId).update(postData)')
bullet('7. Pop screen with result=true')

add_page_break()


# ═══════════════════════════════════════════════════════════
# SCREEN 5 — POST DETAIL SCREEN
# ═══════════════════════════════════════════════════════════
screen_label(5)
doc.add_heading('Post Detail Screen', level=1)
file_line('lib/screens/near by/near_by_post_detail_screen.dart')
purpose_line('Full-detail view of a single post. Shows images (with zoom viewer), title, price/offer chips, category, description, keywords, and provides Chat + Voice Call actions. Adapts headings based on tab category.')

add_screenshot_placeholder('Post Detail Screen')

doc.add_heading('Constructor Parameters', level=2)
add_table(
    ['Parameter', 'Type', 'Required', 'Default', 'Description'],
    [
        ['postId', 'String', 'Yes', '\u2014', 'Firestore document ID'],
        ['post', 'Map<String, dynamic>', 'Yes', '\u2014', 'Full post data map'],
        ['distanceText', 'String', 'No', '""', 'Pre-calculated distance'],
        ['isDeleted', 'bool', 'No', 'false', 'Shows restore/delete buttons'],
        ['tabCategory', 'String', 'No', '"Products"', 'Services, Jobs, Products, Donation'],
    ],
    col_widths=[1417, 1984, 1000, 1000, 3670]
)

doc.add_heading('AppBar (Chat-style)', level=2)
add_table(
    ['Element', 'Position', 'Behavior'],
    [
        ['Back arrow', 'Left', 'iOS-style, pops with _isSaved result'],
        ['Avatar', 'Center-left', 'CircleAvatar: photo or initial letter'],
        ['Username', 'Center', 'Poppins 17px, w600, max 1 line'],
        ['Call indicator', 'Next to name', 'Green circle + phone icon (if allowCalls)'],
        ['Bookmark', 'Right', 'Toggle save/unsave (hidden for deleted)'],
    ],
    col_widths=[1984, 1700, 5387]
)

doc.add_heading('Image Grid Layout', level=2)
add_table(
    ['Images', 'Layout', 'Height'],
    [
        ['1', 'Single full-width image', '280px'],
        ['2', 'Two side-by-side equal images', '280px'],
        ['3', 'One large left (3/5) + two stacked right (2/5)', '280px'],
    ],
    col_widths=[1134, 5103, 2835]
)
bullet('Gradient overlays on all 4 edges + 2 corners (cinematic effect)')
bullet('Tap any image: Full-screen InteractiveViewer (pinch-to-zoom)')
bullet('Swipeable PageView with image counter ("1 / 3")')
bullet('Close button: Top-right circle')

doc.add_heading('Tab-Specific Section Headings', level=2)
add_table(
    ['Tab Category', 'Category Heading', 'Description Heading', 'Highlights Heading'],
    [
        ['Services', 'Service Type', 'Service Details', 'Features'],
        ['Jobs', 'Job Category', 'Job Description', 'Requirements'],
        ['Donation', 'Donation Type', 'About Donation', 'Details'],
        ['Products (default)', 'Category', 'Description', 'Highlights'],
    ],
    col_widths=[1984, 2268, 2268, 2551]
)

doc.add_heading('Tab-Specific Info Chips', level=2)
add_table(
    ['Tab', 'Chip 1', 'Chip 2', 'Chip 3'],
    [
        ['Services', 'Price (currency + amount)', '"Available" (green)', 'Distance'],
        ['Jobs', 'Salary (currency + amount)', 'Job Type (orange)', 'Distance'],
        ['Products', 'Price (currency + amount)', 'Offer or "Best Price" (green)', 'Distance'],
        ['Donation', '"Free" (pink)', '"Donation" (green)', 'Distance'],
    ],
    col_widths=[1417, 2835, 2835, 1984]
)
bullet('Own post: Distance chip shows location name (e.g., "Nagpur")')
bullet('Other user: Distance chip shows distance (e.g., "2.3 km")')

doc.add_heading('Content Sections', level=2)
add_table(
    ['Section', 'Style', 'Details'],
    [
        ['Title', '22px, bold, white', 'Max 3 lines, overflow ellipsis, height 1.2'],
        ['Info Chips', 'Horizontal scrollable row', 'Price + Offer/Status + Distance'],
        ['Category', 'Wrapped chips', 'Blue accent bg, category icon + text'],
        ['Description', 'Glass card', 'Grey text, 14px, line height 1.7'],
        ['Highlights', 'Glass card with chip wrap', 'Check icon + keyword text'],
    ],
    col_widths=[1700, 2551, 4820]
)

doc.add_heading('Save/Bookmark Feature', level=2)
bullet('Checks saved status on init from users/{uid}/saved_posts/{postId}')
bullet('Toggle with haptic feedback')
bullet('Returns _isSaved to parent screen via Navigator.pop for state sync')

doc.add_heading('Bottom Action Bar (Other User\'s Post)', level=2)
add_table(
    ['Button', 'Style', 'Action'],
    [
        ['Chat', 'Blue gradient (#016CFF), rounded 14px', 'Opens EnhancedChatScreen (source: "NearBy")'],
        ['Call', 'Green gradient, rounded 14px', 'Creates call doc + notification + VoiceCallScreen'],
    ],
    col_widths=[1417, 3402, 4252]
)
bullet('Hidden for own posts (isOwnPost check)')
bullet('Call button only shown when allowCalls != false')

doc.add_heading('Bottom Action Bar (Deleted Post)', level=2)
add_table(
    ['Button', 'Style', 'Action'],
    [
        ['Restore', 'Green gradient', 'Sets isActive: true, removes deletedAt, pops screen'],
        ['Delete Forever', 'Red gradient', 'Permanently deletes document, pops screen'],
    ],
    col_widths=[1700, 2268, 5103]
)

doc.add_heading('Voice Call Flow', level=2)
bullet('1. Validates not calling self')
bullet('2. Fetches caller + receiver data from Firestore users collection')
bullet('3. Creates call document in "calls" collection:')
add_code_block('{ callerId, receiverId, callerName, callerPhoto,\n  receiverName, receiverPhoto,\n  participants: [callerId, receiverId],\n  status: "calling", type: "audio",\n  source: "NearBy", timestamp: serverTimestamp() }')
bullet('4. Sends push notification via NotificationService.sendNotificationToUser()')
bullet('5. Navigates to VoiceCallScreen(callId, otherUser, isOutgoing: true)')

doc.add_heading('Chat Flow', level=2)
bullet('1. Validates not chatting with self')
bullet('2. Fetches receiver user data from Firestore')
bullet('3. Builds UserProfile object (uid, name, email, photo, location, coordinates)')
bullet('4. Navigates to EnhancedChatScreen(otherUser, source: "NearBy")')

doc.add_heading('Distance Computation', level=2)
bullet('Gets post coordinates from post["latitude"] / post["longitude"]')
bullet('Gets user location: Firestore user doc first, fallback to GPS last known')
bullet('Haversine formula: dist = R * 2 * atan2(sqrt(a), sqrt(1-a))')
bullet('Display: "X m" (< 1 km) or "X.X km" (>= 1 km)')

add_page_break()


# ═══════════════════════════════════════════════════════════
# SCREEN NAVIGATION FLOW
# ═══════════════════════════════════════════════════════════
doc.add_heading('Screen Navigation Flow', level=1)

add_code_block(
    'MainNavigationScreen (Tab: Nearby)\n'
    '|\n'
    '+-- NearByScreen (Feed with 4 category tabs)\n'
    '    |\n'
    '    +-- [More btn] --> MyPostsScreen\n'
    '    |   +-- My Post tab --> Post cards\n'
    '    |   |   +-- [Edit] --> EditPostScreen\n'
    '    |   |   +-- [Delete] --> Soft delete dialog\n'
    '    |   +-- Saved tab --> Saved post cards\n'
    '    |   |   +-- [Unsave] --> Removes bookmark\n'
    '    |   +-- Delete tab --> Deleted post cards\n'
    '    |       +-- [Restore] --> Restore dialog\n'
    '    |       +-- [Delete Forever] --> Permanent delete\n'
    '    |\n'
    '    +-- [+ FAB] --> CreatePostScreen\n'
    '    |   +-- (On success) --> Returns to NearByScreen\n'
    '    |\n'
    '    +-- [Post Card Tap] --> NearByPostDetailScreen\n'
    '        +-- [Chat] --> EnhancedChatScreen\n'
    '        +-- [Call] --> VoiceCallScreen\n'
    '        +-- [Bookmark] --> Toggle save\n'
    '        +-- [Image Tap] --> Full-screen viewer'
)

add_page_break()


# ═══════════════════════════════════════════════════════════
# FIRESTORE COLLECTIONS USED
# ═══════════════════════════════════════════════════════════
doc.add_heading('Firestore Collections Used', level=1)

add_table(
    ['Collection', 'Usage'],
    [
        ['posts', 'Read (stream): Feed posts ordered by createdAt desc, limit 200'],
        ['posts', 'Read (stream): My Post tab — userId + isActive=true, limit 50'],
        ['posts', 'Read (stream): Delete tab — userId + isActive=false, limit 100'],
        ['posts', 'Write: CreatePostScreen — posts.add(postData)'],
        ['posts', 'Update: EditPostScreen — posts.doc(id).update(postData)'],
        ['posts', 'Update: Soft delete — isActive=false, deletedAt=serverTimestamp'],
        ['posts', 'Update: Restore — isActive=true, remove deletedAt'],
        ['posts', 'Delete: Permanent delete — posts.doc(id).delete()'],
        ['users/{uid}/saved_posts', 'Read/Write/Delete: Save and unsave bookmarks'],
        ['users', 'Read: Fetch user profile data across all screens'],
        ['calls', 'Write: Create call document for voice calling'],
        ['Firebase Storage', 'Write: Upload post images to posts/{userId}/'],
    ],
    col_widths=[2835, 6236]
)

add_page_break()


# ═══════════════════════════════════════════════════════════
# SHARED UI PATTERNS
# ═══════════════════════════════════════════════════════════
doc.add_heading('Shared UI Patterns', level=1)

doc.add_heading('Glassmorphism Components', level=2)
add_table(
    ['Component', 'Usage'],
    [
        ['Post card bottom bar', 'Blur 12, black 55% bg, white 15% border'],
        ['Distance badge', 'Blur 10, black 45% bg, white 25% border, 8px radius'],
        ['Save button', 'Blur 8, blue 85% bg, white 40% border, circle'],
        ['Info glass card', 'Blur 10, white 6% bg, white 30% border, 16px radius'],
        ['Bottom action bar', 'Blur 20, black 50% bg, white 8% top border'],
        ['Confirmation dialogs', 'Gradient #404040 to #0F0F0F, white 15% border, 20px radius'],
    ],
    col_widths=[2835, 6236]
)

doc.add_heading('Floating Card Animation', level=2)
add_table(
    ['Property', 'Value'],
    [
        ['Widget', '_FloatingCard (StatefulWidget + SingleTickerProvider)'],
        ['Vertical range', '-6px to +6px'],
        ['Duration', '1600-2920ms (staggered: 1600 + (index % 6) * 220)'],
        ['Initial offset', '(index * 0.17) % 1.0'],
        ['Curve', 'Curves.easeInOut'],
        ['Repeat', 'reverse: true (bouncing)'],
    ],
    col_widths=[2835, 6236]
)

doc.add_heading('Background Gradient', level=2)
para('All screens: LinearGradient top-to-bottom, Color(64, 64, 64) to Color(0, 0, 0)', size_emu=127000)

doc.add_heading('Typography', level=2)
add_table(
    ['Element', 'Specification'],
    [
        ['Font family', 'Poppins (throughout all screens)'],
        ['Screen titles', '20px, bold, white'],
        ['Section headings', '18px, weight 700, white'],
        ['Body text', '14px, white or grey[400]'],
        ['Prices', '13px, #00D67D (green), bold'],
        ['Badges/chips', '10-13px, white, bold'],
        ['Card username', '13px, white, weight 700'],
    ],
    col_widths=[2835, 6236]
)

doc.add_heading('Animations', level=2)
add_table(
    ['Animation', 'Description'],
    [
        ['FloatingCard', 'Subtle vertical bounce, staggered timing per card index'],
        ['FlipPageRoute', 'Custom 3D flip transition for card->detail navigation'],
        ['Tab switch', '300ms loading indicator on category tab change'],
        ['Haptic feedback', 'lightImpact on taps, mediumImpact on significant actions'],
    ],
    col_widths=[2835, 6236]
)

add_page_break()


# ═══════════════════════════════════════════════════════════
# KEY BUSINESS RULES
# ═══════════════════════════════════════════════════════════
doc.add_heading('Key Business Rules', level=1)

add_table(
    ['#', 'Rule', 'Description'],
    [
        ['1', '5 km radius', 'Only posts within 5 km of user location appear in feed'],
        ['2', 'One post per user', 'Feed shows only the latest post per user (dedup by userId)'],
        ['3', 'Images required', 'Posts without at least one image are hidden from feed'],
        ['4', '30-day auto-delete', 'Soft-deleted posts are permanently removed after 30 days'],
        ['5', '30-day expiry', 'New/edited posts get expiresAt set to 30 days from creation/edit'],
        ['6', 'No self-interaction', 'Users cannot chat with or call themselves'],
        ['7', 'Donation isolation', 'Donation posts only appear in Donation tab, never in others'],
        ['8', 'Call permission', 'Users toggle allowCalls per post to control call availability'],
        ['9', 'Privacy', 'Location shows city name only, not exact GPS coordinates'],
        ['10', 'Max 3 images', 'Create and Edit screens limit to 3 images per post'],
    ],
    col_widths=[567, 1984, 6520]
)


# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
output_path = r'c:\Users\teamp\OneDrive\Documents\Single-Tap\NEARBY_SCREENS_DOCUMENTATION.docx'
doc.save(output_path)
print(f'Done! Saved to {output_path}')
