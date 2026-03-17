#!/usr/bin/env python3
"""
generate_nearby_doc_v3.py
─────────────────────────
Reads NEARBY_SCREENS_DOCUMENTATION.md and converts it into a professionally
styled DOCX using the EXACT same styling as generate_networking_doc.py.
Only the content-specific values (title, module name, paths) differ.
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════════
# CREATE DOCUMENT & APPLY PAGE SETUP  (exact match networking)
# ═══════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.top_margin = 720090
    sec.bottom_margin = 720090
    sec.left_margin = 899795
    sec.right_margin = 899795

# ═══════════════════════════════════════════════════════════
# GLOBAL STYLES  (exact match networking)
# ═══════════════════════════════════════════════════════════
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_before = Emu(0)
style.paragraph_format.space_after = Emu(25400)    # 2pt — tight

# H1: 13pt bold #003366
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Emu(165100)  # 13pt
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
h1.paragraph_format.space_before = Emu(101600)  # 8pt
h1.paragraph_format.space_after = Emu(38100)    # 3pt
h1.paragraph_format.keep_with_next = True
h1.paragraph_format.keep_together = True

# H2: 12pt bold #00508C
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Emu(152400)  # 12pt
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x00, 0x50, 0x8C)
h2.paragraph_format.space_before = Emu(76200)   # 6pt
h2.paragraph_format.space_after = Emu(25400)    # 2pt
h2.paragraph_format.keep_with_next = True
h2.paragraph_format.keep_together = True

# H3: 11pt bold #0064A0
h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Emu(139700)  # 11pt
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x00, 0x64, 0xA0)
h3.paragraph_format.space_before = Emu(50800)   # 4pt
h3.paragraph_format.space_after = Emu(25400)    # 2pt
h3.paragraph_format.keep_with_next = True
h3.paragraph_format.keep_together = True

# ═══════════════════════════════════════════════════════════
# HELPER FUNCTIONS  (exact match networking)
# ═══════════════════════════════════════════════════════════
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def add_page_break():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(0)
    p.paragraph_format.space_after = Emu(0)
    run = p.add_run()
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)


def screen_label(number):
    """SCREEN X — 9pt, bold, #0066B2"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(0)
    run = p.add_run(f'SCREEN {number}')
    run.bold = True
    run.font.size = Emu(114300)  # 9pt
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xB2)
    run.font.name = 'Calibri'


def file_line(path):
    """File: path"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(12700)  # 1pt
    r1 = p.add_run('File: ')
    r1.bold = True
    r1.font.size = Emu(127000)
    r1.font.name = 'Calibri'
    r2 = p.add_run(path)
    r2.font.size = Emu(127000)
    r2.font.color.rgb = RGBColor(0x00, 0x66, 0xB2)
    r2.font.name = 'Calibri'


def purpose_line(text):
    """Purpose: text"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Emu(38100)  # 3pt
    r1 = p.add_run('Purpose: ')
    r1.bold = True
    r1.font.size = Emu(127000)
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Emu(127000)
    r2.font.name = 'Calibri'


def add_spacer(pts=8):
    """Add an empty spacer paragraph with specified height."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(0)
    p.paragraph_format.space_after = Emu(int(pts * 12700))
    r = p.add_run('')
    r.font.size = Emu(12700)  # 1pt invisible text


def para(text, bold=False, size_emu=127000, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Emu(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Emu(size_emu)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Emu(127000)
    run.font.name = 'Calibri'


def bullet_bold(prefix, rest):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.size = Emu(127000)
    r1.font.name = 'Calibri'
    r2 = p.add_run(rest)
    r2.font.size = Emu(127000)
    r2.font.name = 'Calibri'


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_vertical_padding(cell, top_dxa=60, bottom_dxa=60):
    """Add vertical padding to a cell to increase row height."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(f'{{{W_NS}}}tcMar')
    if existing is not None:
        tcPr.remove(existing)
    mar_xml = (
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top_dxa}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom_dxa}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(parse_xml(mar_xml))


def set_cell_text(cell, text, bold=False, color=None, size_emu=127000):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Emu(0)
    p.paragraph_format.space_after = Emu(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Emu(size_emu)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    set_cell_vertical_padding(cell)
    # Vertical center alignment
    tcPr = cell._tc.get_or_add_tcPr()
    existing_vAlign = tcPr.find(f'{{{W_NS}}}vAlign')
    if existing_vAlign is not None:
        tcPr.remove(existing_vAlign)
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    existing = tcPr.find(f'{{{W_NS}}}tcW')
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def keep_table_on_one_page(table):
    """Prevent table from splitting across pages."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
        trPr.append(cant_split)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._element.get_or_add_pPr()
                keep_next = parse_xml(f'<w:keepNext {nsdecls("w")}/>')
                keep_lines = parse_xml(f'<w:keepLines {nsdecls("w")}/>')
                pPr.append(keep_next)
                pPr.append(keep_lines)


def add_table(headers, rows, col_widths=None):
    """Table Grid with #003366 header, white text."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '003366')
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        if col_widths:
            set_cell_width(cell, col_widths[i])
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_text(cell, val)
            if col_widths:
                set_cell_width(cell, col_widths[ci])
    # Only keep small tables together; large tables split to avoid blank gaps
    if len(rows) <= 10:
        keep_table_on_one_page(table)
    return table


def _add_paragraph_border(paragraph, color='999999', style='dashed', size='8'):
    """Add a box border around a paragraph using pBdr XML."""
    pPr = paragraph._element.get_or_add_pPr()
    bdr_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="{style}" w:sz="{size}" w:color="{color}" w:space="8"/>'
        f'  <w:bottom w:val="{style}" w:sz="{size}" w:color="{color}" w:space="8"/>'
        f'  <w:left w:val="{style}" w:sz="{size}" w:color="{color}" w:space="8"/>'
        f'  <w:right w:val="{style}" w:sz="{size}" w:color="{color}" w:space="8"/>'
        f'</w:pBdr>'
    )
    pPr.append(parse_xml(bdr_xml))


def _add_paragraph_shading(paragraph, fill='F5F5F5'):
    """Add background shading to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    pPr.append(shading)


def add_screenshot_placeholder(screen_name):
    """Screenshot placeholder — uses plain paragraphs (NOT a table).

    Plain paragraphs accept image paste in Word without any restrictions.
    The user can click on the placeholder text and paste/insert an image.
    """
    add_spacer(4)

    # Label paragraph — screen name
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_label.paragraph_format.space_before = Emu(0)
    p_label.paragraph_format.space_after = Emu(0)
    r_label = p_label.add_run(f'\U0001F4F1  [ {screen_name} ]')
    r_label.font.size = Emu(139700)  # 11pt
    r_label.bold = True
    r_label.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    r_label.font.name = 'Calibri'

    # Main paste area — single paragraph with border + shading
    # User clicks here, deletes placeholder text, and pastes screenshot
    p_paste = doc.add_paragraph()
    p_paste.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_paste.paragraph_format.space_before = Emu(508000)   # 40pt top padding
    p_paste.paragraph_format.space_after = Emu(508000)    # 40pt bottom padding
    _add_paragraph_border(p_paste)
    _add_paragraph_shading(p_paste)

    r_paste = p_paste.add_run('Click here and paste screenshot (Ctrl+V)')
    r_paste.font.size = Emu(127000)  # 10pt
    r_paste.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r_paste.font.name = 'Calibri'

    add_spacer(4)


def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.3)
    # Grey background shading on the paragraph
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
    pPr.append(shading)


def add_inline_formatted_paragraph(text, is_bullet=False, is_numbered=False):
    """Add a paragraph with inline **bold**, *italic*, `code` formatting."""
    if is_bullet:
        p = doc.add_paragraph(style='List Bullet')
    elif is_numbered:
        p = doc.add_paragraph(style='List Number')
    else:
        p = doc.add_paragraph()

    # Parse inline formatting: **bold**, `code`, *italic* (in that precedence order)
    pattern = r'(\*\*(.+?)\*\*|`([^`]+)`|\*([^*]+?)\*|([^*`]+))'

    for match in re.finditer(pattern, text):
        bold_text = match.group(2)
        code_text = match.group(3)
        italic_text = match.group(4)
        plain_text = match.group(5)

        if bold_text is not None:
            run = p.add_run(bold_text)
            run.bold = True
            run.font.size = Emu(127000)
            run.font.name = 'Calibri'
        elif code_text is not None:
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        elif italic_text is not None:
            run = p.add_run(italic_text)
            run.italic = True
            run.font.size = Emu(127000)
            run.font.name = 'Calibri'
        elif plain_text is not None:
            run = p.add_run(plain_text)
            run.font.size = Emu(127000)
            run.font.name = 'Calibri'


# ═══════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ═══════════════════════════════════════════════════════════

def strip_inline_formatting(text):
    """Remove markdown inline markers for plain text extraction."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\*([^*]+?)\*', r'\1', text)
    return text.strip()


def is_ascii_art(code_block):
    """Detect if a code block is ASCII art (boxes, diagrams) vs actual code."""
    art_chars = set('┌┐└┘├┤┬┴─│┼╔╗╚╝╠╣╦╩═║←→↑↓▪▫●○■□▲△▼▽')
    lines = code_block.strip().split('\n')
    if not lines:
        return False
    art_count = 0
    for line in lines:
        for ch in line:
            if ch in art_chars:
                art_count += 1
    return art_count > 5


def parse_table(lines):
    """Parse a markdown table (list of lines) into headers and rows."""
    if len(lines) < 2:
        return None, None

    def parse_row(line):
        cells = line.strip().strip('|').split('|')
        return [c.strip() for c in cells]

    headers = parse_row(lines[0])

    # lines[1] is the separator (---)
    rows = []
    for line in lines[2:]:
        if line.strip():
            row = parse_row(line)
            while len(row) < len(headers):
                row.append('')
            rows.append(row[:len(headers)])

    return headers, rows


def compute_col_widths(headers, rows):
    """Compute proportional column widths in dxa (total ~9071 dxa = page width)."""
    total_dxa = 9071
    num_cols = len(headers)
    if num_cols == 0:
        return []

    max_lens = []
    for i in range(num_cols):
        col_max = len(strip_inline_formatting(headers[i]))
        for row in rows:
            if i < len(row):
                col_max = max(col_max, len(strip_inline_formatting(row[i])))
        max_lens.append(max(col_max, 3))

    total_len = sum(max_lens)
    widths = []
    for ml in max_lens:
        w = int((ml / total_len) * total_dxa)
        widths.append(max(w, 567))

    diff = total_dxa - sum(widths)
    if widths:
        widths[-1] += diff

    return widths


def extract_screen_info(lines, start_idx):
    """
    Extract File:, Class:, Lines:, Purpose: from the lines immediately
    following a screen heading (## N. Screen Name).
    Returns (file_path, purpose_text, consumed_count).
    """
    file_path = ''
    purpose_text = ''
    consumed = 0
    idx = start_idx

    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            consumed += 1
            continue

        m = re.match(r'\*\*File:\*\*\s*`?([^`]+)`?', line)
        if m:
            file_path = m.group(1).strip()
            idx += 1
            consumed += 1
            continue

        m = re.match(r'\*\*Class:\*\*', line)
        if m:
            idx += 1
            consumed += 1
            continue

        m = re.match(r'\*\*Lines:\*\*', line)
        if m:
            idx += 1
            consumed += 1
            continue

        m = re.match(r'\*\*Purpose:\*\*\s*(.*)', line)
        if m:
            purpose_text = m.group(1).strip()
            idx += 1
            consumed += 1
            continue

        break

    return file_path, purpose_text, consumed


# ═══════════════════════════════════════════════════════════
# READ AND PARSE MARKDOWN
# ═══════════════════════════════════════════════════════════
md_path = r'c:\Users\teamp\OneDrive\Documents\Single-Tap\NEARBY_SCREENS_DOCUMENTATION.md'

with open(md_path, 'r', encoding='utf-8') as f:
    md_content = f.read()

lines = md_content.split('\n')

# ═══════════════════════════════════════════════════════════
# TITLE PAGE  (exact match networking layout)
# ═══════════════════════════════════════════════════════════
for _ in range(2):
    doc.add_paragraph()

# SINGLE TAP — 28pt, bold, #003366
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Emu(25400)  # 2pt
run = p.add_run('SINGLE TAP')
run.bold = True
run.font.size = Emu(355600)  # 28pt
run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
run.font.name = 'Calibri'

# Nearby Module — 22pt, #0066B2
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Emu(25400)  # 2pt
run = p.add_run('Nearby Module')
run.font.size = Emu(279400)  # 22pt
run.font.color.rgb = RGBColor(0x00, 0x66, 0xB2)
run.font.name = 'Calibri'

# Complete Screen Documentation — 14pt, #646464
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Emu(76200)  # 6pt
run = p.add_run('Complete Screen Documentation')
run.font.size = Emu(177800)  # 14pt
run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
run.font.name = 'Calibri'

# Extract metadata from the first few lines of MD
md_module = 'Nearby / Marketplace'
md_total_screens = '3 (1 main feed + 1 saved posts + 1 post detail)'
md_last_updated = '16 Mar 2026'
md_stack = 'Flutter, Firebase, Firestore, HTTP API'

for line in lines[:10]:
    m = re.match(r'\*\*Module:\*\*\s*(.*)', line)
    if m:
        md_module = m.group(1).strip()
    m = re.match(r'\*\*Total Screens:\*\*\s*(.*)', line)
    if m:
        md_total_screens = m.group(1).strip()
    m = re.match(r'\*\*Last Updated:\*\*\s*(.*)', line)
    if m:
        md_last_updated = m.group(1).strip()

# Meta table
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Module', md_module),
    ('Total Screens', md_total_screens),
    ('Stack', md_stack),
    ('Last Updated', md_last_updated),
]
for i, (key, val) in enumerate(meta_data):
    cell_k = meta_table.rows[i].cells[0]
    cell_v = meta_table.rows[i].cells[1]
    set_cell_shading(cell_k, '003366')
    set_cell_text(cell_k, key, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_width(cell_k, 2835)
    set_cell_shading(cell_v, 'F0F5FF')
    set_cell_text(cell_v, val)
    set_cell_width(cell_v, 5669)
keep_table_on_one_page(meta_table)

add_page_break()


# ═══════════════════════════════════════════════════════════
# PROCESS MARKDOWN BODY LINE BY LINE
# ═══════════════════════════════════════════════════════════

screen_counter = 0
screen_heading_pattern = re.compile(r'^##\s+(\d+)\.\s+(.*)')
toc_emitted = False

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # ── Skip the very first lines (title + metadata) that we already handled ──
    if i < 10 and (stripped.startswith('# Nearby Module') or
                    stripped.startswith('**App:**') or
                    stripped.startswith('**Module:**') or
                    stripped.startswith('**Total Screens:**') or
                    stripped.startswith('**Last Updated:**')):
        i += 1
        continue

    # ── Horizontal rule → skip ──
    if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
        i += 1
        continue

    # ── Fenced code block ── ```
    if stripped.startswith('```'):
        lang = stripped[3:].strip()
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing ```

        code_text = '\n'.join(code_lines)
        add_code_block(code_text)
        continue

    # ── Table of Contents (## Table of Contents) ──
    if stripped == '## Table of Contents' and not toc_emitted:
        toc_emitted = True
        doc.add_heading('Table of Contents', level=1)

        toc_entries = []
        i += 1
        while i < len(lines):
            tl = lines[i].strip()
            if not tl:
                i += 1
                continue
            m = re.match(r'(\d+)\.\s+\[(.+?)\]\(#.*?\)', tl)
            if m:
                toc_entries.append((m.group(1), m.group(2)))
                i += 1
            else:
                break

        if toc_entries:
            toc_table = doc.add_table(rows=1 + len(toc_entries), cols=2)
            toc_table.style = 'Table Grid'
            toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, h in enumerate(['#', 'Section']):
                cell = toc_table.rows[0].cells[ci]
                set_cell_shading(cell, '003366')
                set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                set_cell_width(cell, [850, 8221][ci])
            for ri, (num, name) in enumerate(toc_entries):
                set_cell_text(toc_table.rows[ri + 1].cells[0], num)
                set_cell_width(toc_table.rows[ri + 1].cells[0], 850)
                set_cell_text(toc_table.rows[ri + 1].cells[1], name)
                set_cell_width(toc_table.rows[ri + 1].cells[1], 8221)
            keep_table_on_one_page(toc_table)

        add_page_break()
        continue

    # ── Numbered screen section heading: ## N. Screen Name ──
    m = screen_heading_pattern.match(stripped)
    if m:
        screen_num = m.group(1)
        screen_name = m.group(2).strip()
        screen_counter = int(screen_num)

        # Add page break before screens 2+
        if screen_counter > 1:
            add_page_break()

        screen_label(screen_counter)
        doc.add_heading(screen_name, level=1)

        i += 1
        file_path, purpose_text, consumed = extract_screen_info(lines, i)
        i += consumed

        if file_path:
            file_line(file_path)
        if purpose_text:
            purpose_line(purpose_text)

        add_screenshot_placeholder(screen_name)
        continue

    # ── Non-numbered ## heading ──
    m2 = re.match(r'^##\s+(.*)', stripped)
    if m2 and not screen_heading_pattern.match(stripped):
        heading_text = m2.group(1).strip()
        doc.add_heading(heading_text, level=1)
        i += 1
        continue

    # ── ### heading → H2 ──
    m3 = re.match(r'^###\s+(.*)', stripped)
    if m3:
        heading_text = m3.group(1).strip()
        heading_text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', heading_text)
        heading_text = heading_text.replace('`', '')
        doc.add_heading(heading_text, level=2)
        i += 1
        continue

    # ── #### heading → H3 ──
    m4 = re.match(r'^####\s+(.*)', stripped)
    if m4:
        heading_text = m4.group(1).strip()
        heading_text = heading_text.replace('`', '')
        doc.add_heading(heading_text, level=3)
        i += 1
        continue

    # ── Markdown table ──
    if '|' in stripped and not stripped.startswith('```'):
        table_lines = []
        j = i
        while j < len(lines) and '|' in lines[j].strip() and lines[j].strip():
            table_lines.append(lines[j].strip())
            j += 1

        if len(table_lines) >= 2:
            sep_line = table_lines[1]
            if re.match(r'^[\s|:-]+$', sep_line):
                headers, rows = parse_table(table_lines)
                if headers and rows is not None:
                    clean_headers = [strip_inline_formatting(h) for h in headers]
                    clean_rows = [[strip_inline_formatting(c) for c in row] for row in rows]

                    col_widths = compute_col_widths(clean_headers, clean_rows)
                    add_table(clean_headers, clean_rows, col_widths)

                i = j
                continue

    # ── Bullet list item: - text or * text ──
    m_bullet = re.match(r'^[-*]\s+(.*)', stripped)
    if m_bullet:
        bullet_text = m_bullet.group(1).strip()
        m_bp = re.match(r'\*\*(.+?)\*\*\s*(.*)', bullet_text)
        if m_bp:
            prefix = m_bp.group(1)
            rest = m_bp.group(2)
            if rest.startswith(':'):
                bullet_bold(prefix, rest)
            elif prefix.endswith(':'):
                bullet_bold(prefix + ' ', rest)
            else:
                bullet_bold(prefix + ' ', rest)
        else:
            plain = strip_inline_formatting(bullet_text)
            bullet(plain)
        i += 1
        continue

    # ── Numbered list item: N. text ──
    m_num = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if m_num and not screen_heading_pattern.match(stripped):
        num_text = m_num.group(2).strip()
        m_bp = re.match(r'\*\*(.+?)\*\*\s*(.*)', num_text)
        if m_bp:
            prefix = m_bp.group(1)
            rest = m_bp.group(2)
            p = doc.add_paragraph(style='List Number')
            r1 = p.add_run(prefix + ' ')
            r1.bold = True
            r1.font.size = Emu(127000)
            r1.font.name = 'Calibri'
            r2 = p.add_run(strip_inline_formatting(rest))
            r2.font.size = Emu(127000)
            r2.font.name = 'Calibri'
        else:
            plain = strip_inline_formatting(num_text)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(plain)
            run.font.size = Emu(127000)
            run.font.name = 'Calibri'
        i += 1
        continue

    # ── Bold paragraph: **text** (standalone line) ──
    if re.match(r'^\*\*(.+)\*\*$', stripped):
        text = re.match(r'^\*\*(.+)\*\*$', stripped).group(1)
        para(text, bold=True, size_emu=127000)
        i += 1
        continue

    # ── Bold-prefixed line: **label:** rest ──
    m_bold_line = re.match(r'^\*\*(.+?):\*\*\s*(.*)', stripped)
    if m_bold_line:
        label = m_bold_line.group(1).strip()
        rest = m_bold_line.group(2).strip()
        rest = strip_inline_formatting(rest)
        p = doc.add_paragraph()
        r1 = p.add_run(label + ': ')
        r1.bold = True
        r1.font.size = Emu(127000)
        r1.font.name = 'Calibri'
        r2 = p.add_run(rest)
        r2.font.size = Emu(127000)
        r2.font.name = 'Calibri'
        i += 1
        continue

    # ── Indented sub-bullet: starts with spaces + - ──
    m_sub = re.match(r'^(\s{2,})[-*]\s+(.*)', line)
    if m_sub:
        sub_text = strip_inline_formatting(m_sub.group(2).strip())
        try:
            p = doc.add_paragraph(style='List Bullet 2')
        except KeyError:
            p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(sub_text)
        run.font.size = Emu(127000)
        run.font.name = 'Calibri'
        i += 1
        continue

    # ── Non-empty text line (regular paragraph) ──
    if stripped and not stripped.startswith('#'):
        plain = strip_inline_formatting(stripped)
        if plain:
            add_inline_formatted_paragraph(stripped)
        i += 1
        continue

    # ── Empty line or unhandled → skip ──
    i += 1


# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
output_path = r'c:\Users\teamp\OneDrive\Documents\Single-Tap\NEARBY_SCREENS_DOCUMENTATION.docx'
doc.save(output_path)
print(f'Done! Saved to {output_path}')
