"""
Generic Markdown → DOCX converter
Uses same styling as generate_nearby_doc.py (Calibri, blue headings, striped tables)
Usage: python md_to_docx.py <input.md> <output.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════════
# SETUP
# ═══════════════════════════════════════════════════════════
if len(sys.argv) < 3:
    print("Usage: python md_to_docx.py <input.md> <output.docx>")
    sys.exit(1)

md_path = sys.argv[1]
docx_path = sys.argv[2]

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()

# ═══════════════════════════════════════════════════════════
# PAGE SETUP
# ═══════════════════════════════════════════════════════════
for sec in doc.sections:
    sec.top_margin = 720090
    sec.bottom_margin = 720090
    sec.left_margin = 899795
    sec.right_margin = 899795

# ═══════════════════════════════════════════════════════════
# GLOBAL STYLES
# ═══════════════════════════════════════════════════════════
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(4)
h1.paragraph_format.keep_with_next = True

h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x00, 0x50, 0x8C)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(3)
h2.paragraph_format.keep_with_next = True

h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(11)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x00, 0x64, 0xA0)
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after = Pt(2)
h3.paragraph_format.keep_with_next = True

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ═══════════════════════════════════════════════════════════
# HELPER: add inline formatting (bold, code, italic)
# ═══════════════════════════════════════════════════════════
def add_formatted_runs(paragraph, text, base_size=Pt(10), base_color=None):
    """Parse inline markdown (**bold**, `code`, *italic*) and add runs."""
    # Pattern: **bold** or `code` or *italic*
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*)')
    last = 0
    for m in pattern.finditer(text):
        # Add plain text before match
        if m.start() > last:
            r = paragraph.add_run(text[last:m.start()])
            r.font.name = 'Calibri'
            r.font.size = base_size
            if base_color:
                r.font.color.rgb = base_color
        if m.group(2):  # **bold**
            r = paragraph.add_run(m.group(2))
            r.bold = True
            r.font.name = 'Calibri'
            r.font.size = base_size
            if base_color:
                r.font.color.rgb = base_color
        elif m.group(3):  # `code`
            r = paragraph.add_run(m.group(3))
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        elif m.group(4):  # *italic*
            r = paragraph.add_run(m.group(4))
            r.italic = True
            r.font.name = 'Calibri'
            r.font.size = base_size
            if base_color:
                r.font.color.rgb = base_color
        last = m.end()
    # Remaining text
    if last < len(text):
        r = paragraph.add_run(text[last:])
        r.font.name = 'Calibri'
        r.font.size = base_size
        if base_color:
            r.font.color.rgb = base_color

def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell):
    """Set thin borders on cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)

# ═══════════════════════════════════════════════════════════
# PARSE MARKDOWN LINE BY LINE
# ═══════════════════════════════════════════════════════════
i = 0
total = len(lines)

while i < total:
    line = lines[i].rstrip('\n')

    # --- Horizontal rule → page break (skip if at very top)
    if re.match(r'^-{3,}$', line.strip()):
        if i > 2:  # Don't break on first ---
            p = doc.add_paragraph()
            run = p.add_run()
            br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
            run._element.append(br)
        i += 1
        continue

    # --- Empty line
    if line.strip() == '':
        i += 1
        continue

    # --- Code block (```)
    if line.strip().startswith('```'):
        code_lines = []
        i += 1
        while i < total and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        i += 1  # skip closing ```

        # Add as formatted code block
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # Add shading to paragraph
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        pPr.append(shading)
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        continue

    # --- Table (| col | col |)
    if '|' in line and line.strip().startswith('|'):
        table_lines = []
        while i < total and '|' in lines[i] and lines[i].strip().startswith('|'):
            raw = lines[i].strip()
            # Skip separator rows (|---|---|)
            if re.match(r'^\|[\s\-:|]+\|$', raw):
                i += 1
                continue
            # Parse cells
            cells = [c.strip() for c in raw.split('|')[1:-1]]
            table_lines.append(cells)
            i += 1

        if not table_lines:
            continue

        # Determine column count from header
        col_count = len(table_lines[0])

        # Create table
        table = doc.add_table(rows=len(table_lines), cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True

        # Set table style with borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        for row_idx, cells in enumerate(table_lines):
            row = table.rows[row_idx]
            for col_idx in range(min(len(cells), col_count)):
                cell = row.cells[col_idx]
                cell.text = ''  # Clear default
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

                cell_text = cells[col_idx] if col_idx < len(cells) else ''
                set_cell_borders(cell)

                if row_idx == 0:
                    # Header row: bold, dark blue background
                    shade_cell(cell, '003366')
                    add_formatted_runs(p, cell_text, Pt(9), RGBColor(0xFF, 0xFF, 0xFF))
                    for run in p.runs:
                        run.bold = True
                else:
                    # Alternate row shading
                    if row_idx % 2 == 0:
                        shade_cell(cell, 'F2F7FC')
                    add_formatted_runs(p, cell_text, Pt(9))

        # Add spacing after table
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        continue

    # --- Headings
    heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2).strip()
        if level == 1:
            p = doc.add_heading(text, level=1)
        elif level == 2:
            p = doc.add_heading(text, level=2)
        elif level >= 3:
            p = doc.add_heading(text, level=3)
        i += 1
        continue

    # --- Bullet list (- item or * item)
    bullet_match = re.match(r'^(\s*)[*\-]\s+(.*)', line)
    if bullet_match:
        indent = len(bullet_match.group(1))
        text = bullet_match.group(2)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        level = min(indent // 2, 2)
        left_indent = Pt(18 + level * 18)
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = Pt(-12)
        # Add bullet character
        r = p.add_run('• ')
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        add_formatted_runs(p, text)
        i += 1
        continue

    # --- Numbered list
    num_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
    if num_match:
        indent = len(num_match.group(1))
        num = num_match.group(2)
        text = num_match.group(3)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        level = min(indent // 2, 2)
        left_indent = Pt(18 + level * 18)
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = Pt(-14)
        r = p.add_run(f'{num}. ')
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x00, 0x50, 0x8C)
        add_formatted_runs(p, text)
        i += 1
        continue

    # --- Regular paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_formatted_runs(p, line)
    i += 1

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
doc.save(docx_path)
print(f'Done — saved to {docx_path}')
