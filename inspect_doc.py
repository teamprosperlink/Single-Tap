import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Emu, Cm, Inches
from docx.oxml.ns import qn

doc = Document(r'c:\Users\teamp\OneDrive\Documents\Single-Tap\NETWORKING_SCREENS_DOCUMENTATION.docx')

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# === PAGE MARGINS ===
print("=== PAGE SETUP ===")
for i, sec in enumerate(doc.sections):
    print(f"  top={sec.top_margin} bottom={sec.bottom_margin} left={sec.left_margin} right={sec.right_margin}")
    print(f"  page_width={sec.page_width} page_height={sec.page_height}")

# === PARAGRAPH SPACING ===
print("\n=== PARAGRAPH FORMATS (first 30 non-empty) ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip():
        continue
    pf = p.paragraph_format
    text = p.text[:50].encode('ascii','replace').decode()
    style = p.style.name if p.style else 'None'
    print(f"  P[{i}] style={style} space_before={pf.space_before} space_after={pf.space_after} line_spacing={pf.line_spacing} left_indent={pf.left_indent}")
    count += 1
    if count >= 30:
        break

# === TABLE DIMENSIONS ===
print("\n=== TABLE DIMENSIONS ===")
for i, table in enumerate(doc.tables):
    tbl = table._tbl
    nrows = len(table.rows)
    ncols = len(table.columns)
    style = table.style.name if table.style else 'None'

    # Table width
    tblPr = tbl.tblPr
    tblW = tblPr.find(f'{{{W}}}tblW') if tblPr is not None else None
    tw_val = tblW.get(f'{{{W}}}w') if tblW is not None else 'N/A'
    tw_type = tblW.get(f'{{{W}}}type') if tblW is not None else 'N/A'

    # Row heights
    row_heights = []
    for row in table.rows:
        trPr = row._tr.find(f'{{{W}}}trPr')
        if trPr is not None:
            trH = trPr.find(f'{{{W}}}trHeight')
            if trH is not None:
                h = trH.get(f'{{{W}}}val')
                rule = trH.get(f'{{{W}}}hRule')
                row_heights.append(f"{h}({rule})")
            else:
                row_heights.append('auto')
        else:
            row_heights.append('auto')

    # Column widths
    col_widths = []
    if nrows > 0:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                tcW = tcPr.find(f'{{{W}}}tcW')
                if tcW is not None:
                    cw = tcW.get(f'{{{W}}}w')
                    ct = tcW.get(f'{{{W}}}type')
                    col_widths.append(f"{cw}({ct})")
                else:
                    col_widths.append('auto')
            else:
                col_widths.append('auto')

    # Cell padding
    cell_margins = 'N/A'
    if tblPr is not None:
        tblCellMar = tblPr.find(f'{{{W}}}tblCellMar')
        if tblCellMar is not None:
            parts = []
            for side in ['top', 'bottom', 'start', 'end', 'left', 'right']:
                el = tblCellMar.find(f'{{{W}}}{side}')
                if el is not None:
                    parts.append(f"{side}={el.get(f'{{{W}}}w')}")
            cell_margins = ', '.join(parts) if parts else 'default'

    header_text = table.rows[0].cells[0].text[:30].encode('ascii','replace').decode() if nrows > 0 else ''

    print(f"\n  Table {i}: {nrows}x{ncols} style='{style}'")
    print(f"    header[0]='{header_text}'")
    print(f"    table_width={tw_val} ({tw_type})")
    print(f"    col_widths={col_widths}")
    print(f"    row_heights={row_heights[:5]}")
    print(f"    cell_margins={cell_margins}")

    # Check for screenshot tables (1x1) — get full cell formatting
    if ncols == 1 and nrows == 1:
        cell = table.rows[0].cells[0]
        # Cell vertical alignment
        tc = cell._tc
        tcPr = tc.tcPr
        vAlign = 'N/A'
        if tcPr is not None:
            va = tcPr.find(f'{{{W}}}vAlign')
            if va is not None:
                vAlign = va.get(f'{{{W}}}val')

        # Cell height from row
        print(f"    SCREENSHOT CARD: vAlign={vAlign}")
        # Get all paragraph formats in cell
        for pi, cp in enumerate(cell.paragraphs):
            runs_text = ''.join(r.text for r in cp.runs)[:40].encode('ascii','replace').decode()
            align = cp.alignment
            for r in cp.runs:
                print(f"      para[{pi}]: align={align} text='{r.text[:20].encode('ascii','replace').decode()}' size={r.font.size} bold={r.bold}")
                break

    if i >= 10:
        print("  ... (more tables, stopping)")
        break

# === TITLE PAGE TEXT SIZES ===
print("\n=== TITLE PAGE DETAILS ===")
for i in range(15):
    p = doc.paragraphs[i]
    if not p.text.strip():
        continue
    for r in p.runs:
        text = r.text[:40].encode('ascii','replace').decode()
        print(f"  P[{i}]: text='{text}' size={r.font.size} (={Pt(r.font.size/12700) if r.font.size else 'N/A'}) bold={r.bold} color={r.font.color.rgb if r.font.color and r.font.color.rgb else 'N/A'}")

# === HEADING SIZES (from style definition) ===
print("\n=== HEADING STYLE DEFINITIONS ===")
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    print(f"  H{level}: size={hs.font.size} bold={hs.font.bold} color={hs.font.color.rgb if hs.font.color and hs.font.color.rgb else 'N/A'} name={hs.font.name}")
    # Paragraph format
    pf = hs.paragraph_format
    print(f"       space_before={pf.space_before} space_after={pf.space_after}")
